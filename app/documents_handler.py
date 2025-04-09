from docx import Document
from app.crud import get_empleado_completo
import os

def generar_acta_empleado(empleado_id: int, output_path: str = "generated_docs"):
    data = get_empleado_completo(empleado_id)  # Función que reúne todo lo necesario

    if not data:
        return None

    doc = Document()
    doc.add_heading("Acta de Entrega de Activos", 0)

    doc.add_paragraph(f"Fecha: {data['empleado']['Fecha']}")
    doc.add_paragraph(f"Nombre: {data['empleado']['Nombre']}")
    doc.add_paragraph(f"Identificación: {data['empleado']['Identificacion']}")
    doc.add_paragraph(f"Cargo: {data['empleado']['Cargo']}")
    doc.add_paragraph(f"Dependencia: {data['empleado']['Dependencia']}")
    doc.add_paragraph(f"Ubicación: {data['empleado']['UbicacionOficina']}")

    doc.add_heading("Activos Hardware", level=1)
    for hw in data["hardware"]:
        doc.add_paragraph(f"- {hw['TipoHardware']} ({hw['Marca']}, {hw['Modelo']}) - Serial: {hw['Serial']}")

    doc.add_heading("Licencias Software", level=1)
    for lic in data["licencias"]:
        doc.add_paragraph(f"- {lic['NombreLicencia']} | Usuario: {lic['Usuario']}")

    doc.add_heading("Accesos Web", level=1)
    for acc in data["accesos"]:
        doc.add_paragraph(f"- {acc['URL']} | Usuario: {acc['Usuario']}")

    os.makedirs(output_path, exist_ok=True)
    filepath = os.path.join(output_path, f"Acta_{data['empleado']['Nombre']}.docx")
    doc.save(filepath)
    return filepath